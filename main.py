
from flask import Flask, render_template, request
import boto3
import os
import pandas
from pandas import *
from twilio.rest import Client
from docx import Document
from docx.shared import Pt
import aspose.words as aw
from datetime import *

sagemaker_runtime = boto3.client('sagemaker-runtime',region_name='us-east-1',aws_access_key_id=os.getenv('aws_access_key_id'),aws_secret_access_key=os.getenv('aws_secret_access_key'))
client = Client('AC4c984574c2d7f8a03b6f14092de12fb7','3294be0b222f0048a6f74609d14001cc')
symp_db = [["high_fever", "breathlessness", "mood_swings", "weakness_in_limbs", "cough", "itching", "weight_loss", "indigestion", "fatigue", "stiff_neck", "muscle_wasting", "sunken_eyes", "weight_gain", "bladder_discomfort", "skin_rash", "stomach_pain", "cramps", "yellowish_skin", "joint_pain", "patches_in_throat", "constipation", "pus_filled_pimples", "continuous_sneezing", "back_pain", "burning_micturition", "headache", "chills", "pain_during_bowel_movements", "chest_pain", "acidity", "muscle_weakness", "vomiting", "neck_pain", "shivering"], ["high_fever", "breathlessness", "mood_swings", "weakness_in_limbs", "cough", "foul_smell_ofurine", "sweating", "abdominal_pain", "anxiety", "weight_loss", "knee_pain", "indigestion", "weakness_of_one_body_side", "restlessness", "blackheads", "fatigue", "stiff_neck", "bruising", "sunken_eyes", "dizziness", "weight_gain", "skin_rash", "stomach_pain", "loss_of_appetite", "nausea", "bladder_discomfort", "cramps", "yellowish_skin", "joint_pain", "patches_in_throat", "lethargy", "pus_filled_pimples", "swelling_joints", "nodal_skin_eruptions", "dehydration", "headache", "blister", "pain_in_anal_region", "chills", "pain_during_bowel_movements", "chest_pain", "acidity", "skin_peeling", "cold_hands_and_feets", "vomiting", "ulcers_on_tongue", "neck_pain", "shivering"], ["high_fever", "breathlessness", "mood_swings", "cough", "foul_smell_ofurine", "sweating", "abdominal_pain", "anxiety", "weight_loss", "knee_pain", "weakness_of_one_body_side", "blurred_and_distorted_vision", "restlessness", "swelling_of_stomach", "blackheads", "fatigue", "continuous_feel_of_urine", "bruising", "dizziness", "loss_of_appetite", "stomach_pain", "nausea", "loss_of_balance", "altered_sensorium", "scurring", "yellowish_skin", "joint_pain", "lethargy", "red_sore_around_nose", "movement_stiffness", "bloody_stool", "swelling_joints", "diarrhoea", "dark_urine", "burning_micturition", "nodal_skin_eruptions", "dehydration", "headache", "blister", "spinning_movements", "pain_in_anal_region", "chills", "watering_from_eyes", "obesity", "hip_joint_pain", "chest_pain", "silver_like_dusting", "extra_marital_contacts", "dischromic_patches", "skin_peeling", "cold_hands_and_feets", "vomiting", "ulcers_on_tongue", "neck_pain"], ["high_fever", "breathlessness", "mood_swings", "cough", "swollen_legs", "sweating", "abdominal_pain", "weight_loss", "irritation_in_anus", "blurred_and_distorted_vision", "swelling_of_stomach", "excessive_hunger", "restlessness", "fatigue", "continuous_feel_of_urine", "dizziness", "loss_of_appetite", "loss_of_balance", "nausea", "altered_sensorium", "scurring", "yellowish_skin", "family_history", "lethargy", "red_sore_around_nose", "spotting_urination", "passage_of_gases", "movement_stiffness", "lack_of_concentration", "bloody_stool", "swelling_joints", "yellow_crust_ooze", "diarrhoea", "dark_urine", "burning_micturition", "headache", "spinning_movements", "watering_from_eyes", "distention_of_abdomen", "obesity", "chest_pain", "hip_joint_pain", "silver_like_dusting", "painful_walking", "small_dents_in_nails", "extra_marital_contacts", "dischromic_patches", "vomiting", "yellowing_of_eyes", "irregular_sugar_level"], ["high_fever", "breathlessness", "cough", "swollen_legs", "sweating", "abdominal_pain", "irritation_in_anus", "blurred_and_distorted_vision", "excessive_hunger", "fatigue", "stiff_neck", "inflammatory_nails", "dizziness", "loss_of_appetite", "loss_of_balance", "nausea", "swollen_blood_vessels", "yellowish_skin", "family_history", "lethargy", "spotting_urination", "passage_of_gases", "internal_itching", "lack_of_concentration", "mucoid_sputum", "swelling_joints", "yellow_crust_ooze", "diarrhoea", "dark_urine", "headache", "unsteadiness", "distention_of_abdomen", "chest_pain", "history_of_alcohol_consumption", "painful_walking", "small_dents_in_nails", "yellowing_of_eyes", "irregular_sugar_level"], ["high_fever", "swelled_lymph_nodes", "breathlessness", "sweating", "abdominal_pain", "blurred_and_distorted_vision", "stiff_neck", "prominent_veins_on_calf", "inflammatory_nails", "dizziness", "loss_of_appetite", "nausea", "swollen_blood_vessels", "yellowish_skin", "family_history", "constipation", "puffy_face_and_eyes", "internal_itching", "fast_heart_rate", "mucoid_sputum", "diarrhoea", "dark_urine", "headache", "depression", "malaise", "fluid_overload", "unsteadiness", "obesity", "chest_pain", "history_of_alcohol_consumption", "painful_walking", "yellowing_of_eyes"], ["swelled_lymph_nodes", "breathlessness", "mild_fever", "sweating", "abdominal_pain", "blurred_and_distorted_vision", "excessive_hunger", "prominent_veins_on_calf", "phlegm", "loss_of_appetite", "nausea", "constipation", "puffy_face_and_eyes", "muscle_pain", "fast_heart_rate", "diarrhoea", "dark_urine", "headache", "depression", "malaise", "yellow_urine", "fluid_overload", "obesity", "enlarged_thyroid", "irritability", "yellowing_of_eyes"], ["swelled_lymph_nodes", "mild_fever", "sweating", "abdominal_pain", "excessive_hunger", "phlegm", "loss_of_appetite", "nausea", "muscle_pain", "drying_and_tingling_lips", "diarrhoea", "increased_appetite", "malaise", "yellow_urine", "visual_disturbances", "chest_pain", "enlarged_thyroid", "brittle_nails", "muscle_weakness", "irritability", "yellowing_of_eyes"], ["swelled_lymph_nodes", "mild_fever", "abdominal_pain", "throat_irritation", "polyuria", "phlegm", "loss_of_appetite", "slurred_speech", "pain_behind_the_eyes", "fast_heart_rate", "drying_and_tingling_lips", "diarrhoea", "swollen_extremeties", "increased_appetite", "malaise", "visual_disturbances", "chest_pain", "brittle_nails", "muscle_weakness", "irritability", "yellowing_of_eyes", "toxic_look_(typhos)"], ["mild_fever", "redness_of_eyes", "throat_irritation", "polyuria", "belly_pain", "abnormal_menstruation", "rusty_sputum", "slurred_speech", "pain_behind_the_eyes", "muscle_pain", "fast_heart_rate", "back_pain", "receiving_blood_transfusion", "acute_liver_failure", "swollen_extremeties", "depression", "malaise", "red_spots_over_body", "irritability", "yellowing_of_eyes", "toxic_look_(typhos)"], ["swelled_lymph_nodes", "redness_of_eyes", "sinus_pressure", "belly_pain", "abnormal_menstruation", "coma", "rusty_sputum", "receiving_unsterile_injections", "muscle_pain", "back_pain", "receiving_blood_transfusion", "acute_liver_failure", "palpitations", "depression", "malaise", "red_spots_over_body", "irritability", "yellowing_of_eyes"], ["swelled_lymph_nodes", "stomach_bleeding", "malaise", "muscle_pain", "irritability", "runny_nose", "palpitations", "sinus_pressure", "coma", "abnormal_menstruation", "receiving_unsterile_injections"], ["stomach_bleeding", "malaise", "muscle_pain", "congestion", "red_spots_over_body", "runny_nose", "phlegm", "abnormal_menstruation"], ["congestion", "red_spots_over_body", "chest_pain", "phlegm"], ["chest_pain", "loss_of_smell", "blood_in_sputum"]]
endpoint_name = 'med-receipt-1'
df = read_csv('./datasets/Symptom-severity.csv')
symptoms = df['Symptom'].tolist()
for x in range(len(symptoms)):
	s = symptoms[x]
	s = s.replace('_',' ')
	symptoms[x] = s
app = Flask(__name__)
@app.route('/')
def index():
		df = read_csv('./datasets/Symptom-severity.csv')
		symptoms = df['Symptom'].tolist()
		for x in range(len(symptoms)):
			s = symptoms[x]
			s = s.replace('_',' ')
			symptoms[x] = s
		submitted = False
		lenght = len(symptoms)
		return render_template('index.html', symptoms=symptoms, lenght=lenght)

@app.route('/api/forwardtomobile')
def mobile():
	twilbody = request.url.split('?')[1]
	twilbody = twilbody.replace('+',' ')
	twilbody = twilbody.replace('[','')
	twilbody = twilbody.replace('"','')
	twilbody = twilbody.replace(':','')
	twilbody = twilbody.replace('=', '')
	twilbody = twilbody.replace('%5D','')
	twilbody = twilbody.replace('%22','')
	twilbody = twilbody.replace('%5B','')
	twilbody = twilbody.replace('%3A','')
	twilbody = twilbody.replace('%2F',' ')
	twilbody = twilbody.replace('/',' ')
	twilbody = twilbody.replace('SUGGEST','SUGGESTED')
	twilbody = twilbody.replace(',','')
	twilbody = twilbody.replace("'",'')
	print(twilbody)
	message = client.messages.create(to="+12065568873",from_="+17168991476",body=twilbody)
	return "Your Message Has Been Sent!"
@app.route('/api/model/', methods=('GET', 'POST'))
def model():
	global receipt
	takens = [False,False,False,False,False,False,False,False,False,False,False,False,False,False,False]
	print(takens)
	things = [',',',',',',',',',',',',',',',',',',',',',',',',',',',',',']
	#print(takens)
	dones = []
	for symptom in [x.replace('+','_').replace('=on','') for x in request.url.split('?')[1].split('&')]:
		print(symptom)
		for x in range(len(symp_db)):
			if symptom in symp_db[x]:
				if takens[x] == False:
					if symptom in dones:
						pass
					else:
						takens[x] = True
						print(x)
						things[x] = symptom+','
						dones.append(symptom)
				else:
					print('Inaccuracy detected.')
					print(symptom,takens,dones)
	print(''.join(things))
	response = sagemaker_runtime.invoke_endpoint(EndpointName=endpoint_name,Body=''.join(things),ContentType='text/csv')
	diagnosis = response['Body'].read().decode()
	df_desc = read_csv('datasets/symptom_Description.csv')
	size = len(diagnosis)
	diagnosis = diagnosis[:size - 1]
	diagnosis_desc = None
	for index, row in df_desc.iterrows():
		d=row.to_dict()
		if str(d['Disease']) == diagnosis:
			diagnosis_desc = d['Description']
	dfp = read_csv('datasets/symptom_precaution.csv')
	precautions = []
	for index, row in dfp.iterrows():
		d = row.to_dict()
		if str(d['Disease']) == diagnosis:
			precautions.append(d['Precaution_1'])
			precautions.append(d['Precaution_2'])
			precautions.append(d['Precaution_3'])
			precautions.append(d['Precaution_4'])
	return_list = [diagnosis, diagnosis_desc, precautions]
	document = Document(open('static/Invoice (1).docx','rb'))
	style = document.styles['Normal']
	font = style.font
	font.name = 'Courier New'
	font.size = Pt(16)
	document.paragraphs[3].text = str(dones)
	document.paragraphs[4].text = ""
	document.paragraphs[5].text = ""
	document.paragraphs[8].text = diagnosis
	document.paragraphs[11].text = "9/24/2022"
	document.paragraphs[14].text = str(precautions)
	receipt = []
	for paragraph in document.paragraphs:
		print (paragraph.text)
		receipt.append(paragraph.text)
	# document.save('edited.docx')
	
	# doc = aw.Document("edited.docx")
	# doc.extract_pages(0,1).save("page1.png")
	
	return render_template('app.html', return_list = return_list, symptoms=symptoms, receipt = receipt)
app.run(host='0.0.0.0', port=81, debug=True)

